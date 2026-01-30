"""
Export Handlers for Multiple File Formats
Generates PDF, Excel, Word, and TXT exports from user stories
"""
import json
from pathlib import Path
from datetime import datetime
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from config import Config

class ExportHandler:
    """Handle exports to multiple file formats"""
    
    def __init__(self):
        self.exports_dir = Config.EXPORTS_DIR
    
    def export_to_pdf(self, output_structure: dict, filename: str = None) -> str:
        """
        Export user stories to PDF format
        
        Args:
            output_structure: Result from Task 6 (output transformation)
            filename: Optional custom filename
            
        Returns:
            Path to generated PDF file
        """
        if filename is None:
            filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        filepath = self.exports_dir / filename
        
        # Get PDF structure
        pdf_data = output_structure.get('pdf_structure', {})
        
        # Create PDF
        doc = SimpleDocTemplate(str(filepath), pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)
        
        # Container for PDF elements
        elements = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1F1F1F'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#FFD700'),
            spaceAfter=12
        )
        
        # Title
        title = pdf_data.get('title', 'User Stories')
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Executive Summary
        if 'executive_summary' in pdf_data:
            elements.append(Paragraph('Executive Summary', heading_style))
            elements.append(Paragraph(pdf_data['executive_summary'], styles['Normal']))
            elements.append(Spacer(1, 0.3 * inch))
        
        # Sections (User Stories)
        for section in pdf_data.get('sections', []):
            elements.append(Paragraph(section.get('section_title', ''), heading_style))
            elements.append(Spacer(1, 0.1 * inch))
            
            for story in section.get('stories', []):
                # Story table
                story_data = [
                    ['Story ID', story.get('story_id', '')],
                    ['Title', story.get('title', '')],
                    ['Priority', story.get('priority', '')],
                    ['User Story', story.get('user_story', '')],
                    ['Story Points', story.get('story_points', '')],
                    ['Dependencies', story.get('dependencies', 'None')]
                ]
                
                # Add acceptance criteria
                ac_text = '\n'.join([f"• {ac}" for ac in story.get('acceptance_criteria', [])])
                story_data.append(['Acceptance Criteria', ac_text])
                
                t = Table(story_data, colWidths=[1.5*inch, 5*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#FFFEF0')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1F1F1F')),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                elements.append(t)
                elements.append(Spacer(1, 0.2 * inch))
            
            elements.append(Spacer(1, 0.2 * inch))
        
        # Build PDF
        doc.build(elements)
        
        return str(filepath)
    
    def export_to_excel(self, output_structure: dict, filename: str = None) -> str:
        """
        Export user stories to Excel format (Jira-importable)
        
        Args:
            output_structure: Result from Task 6
            filename: Optional custom filename
            
        Returns:
            Path to generated Excel file
        """
        if filename is None:
            filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        filepath = self.exports_dir / filename
        
        # Get Excel structure
        excel_data = output_structure.get('excel_structure', {})
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = excel_data.get('sheet_name', 'User Stories')
        
        # Headers
        headers = excel_data.get('headers', [])
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color='FFD700', end_color='FFD700', fill_type='solid')
        header_font = Font(bold=True, color='1F1F1F')
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Add data rows
        for row_data in excel_data.get('rows', []):
            row_values = [row_data.get(header, '') for header in headers]
            ws.append(row_values)
        
        # Adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save workbook
        wb.save(filepath)
        
        return str(filepath)
    
    def export_to_word(self, output_structure: dict, filename: str = None) -> str:
        """
        Export user stories to Word format
        
        Args:
            output_structure: Result from Task 6
            filename: Optional custom filename
            
        Returns:
            Path to generated Word file
        """
        if filename is None:
            filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = self.exports_dir / filename
        
        # Get Word structure
        word_data = output_structure.get('word_structure', {})
        
        # Create document
        doc = Document()
        
        # Set document properties
        core_properties = doc.core_properties
        core_properties.author = Config.APP_BRAND
        core_properties.title = word_data.get('document_title', 'User Stories')
        
        # Title
        title = doc.add_heading(word_data.get('document_title', 'User Stories'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process sections
        for section in word_data.get('sections', []):
            level = section.get('heading_level', 1)
            text = section.get('heading_text', '')
            
            doc.add_heading(text, level)
            
            if 'content' in section:
                doc.add_paragraph(section['content'])
            
            # Handle subsections (for stories)
            if 'subsections' in section:
                for subsection in section['subsections']:
                    doc.add_heading(subsection.get('heading_text', ''), subsection.get('heading_level', 2))
                    
                    for story in subsection.get('stories', []):
                        # Add story details
                        story_id = story.get('story_id', '')
                        content = story.get('content', '')
                        
                        p = doc.add_paragraph()
                        p.add_run(f"{story_id}: ").bold = True
                        p.add_run(content)
                        doc.add_paragraph()
            
            # Handle tables
            if 'table' in section:
                table_data = section['table']
                table_headers = table_data.get('headers', [])
                table_rows = table_data.get('rows', [])
                
                table = doc.add_table(rows=1, cols=len(table_headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(table_headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
                # Add rows
                for row in table_rows:
                    row_cells = table.add_row().cells
                    for i, cell_value in enumerate(row):
                        row_cells[i].text = str(cell_value)
        
        # Save document
        doc.save(filepath)
        
        return str(filepath)
    
    def export_to_txt(self, output_structure: dict, filename: str = None) -> str:
        """
        Export user stories to plain text format
        
        Args:
            output_structure: Result from Task 6
            filename: Optional custom filename
            
        Returns:
            Path to generated TXT file
        """
        if filename is None:
            filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        filepath = self.exports_dir / filename
        
        # Get TXT structure
        txt_data = output_structure.get('txt_structure', {})
        
        # Build text content
        content = []
        content.append(txt_data.get('header', ''))
        
        for section in txt_data.get('sections', []):
            content.append('\n' + '=' * 60)
            content.append(section.get('section_title', ''))
            content.append('=' * 60 + '\n')
            
            for story in section.get('stories', []):
                content.append(story.get('story_block', ''))
        
        # Write to file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        return str(filepath)
    
    def export_all_formats(self, output_structure: dict, base_filename: str = None) -> dict:
        """
        Export to all formats at once
        
        Args:
            output_structure: Result from Task 6
            base_filename: Base name for files (without extension)
            
        Returns:
            Dictionary with paths to all generated files
        """
        if base_filename is None:
            base_filename = f"user_stories_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        results = {
            'pdf': self.export_to_pdf(output_structure, f"{base_filename}.pdf"),
            'excel': self.export_to_excel(output_structure, f"{base_filename}.xlsx"),
            'word': self.export_to_word(output_structure, f"{base_filename}.docx"),
            'txt': self.export_to_txt(output_structure, f"{base_filename}.txt")
        }
        
        return results
